"""
Genera ficheros sintéticos de entrada (.docx, .xlsx, .pdf) con el mismo
contenido que ejemplos/papel_trabajo_compras.md, para probar los lectores
sin depender de una exportación real de Pentana.

    .venv/bin/python scripts/generar_ejemplos_entrada.py [carpeta_destino]
"""
import sys
from pathlib import Path

RAIZ = Path(__file__).resolve().parent.parent
DESTINO = Path(sys.argv[1]) if len(sys.argv) > 1 else RAIZ / "ejemplos" / "entrada_sintetica"
DESTINO.mkdir(parents=True, exist_ok=True)

TITULO = "Papel de trabajo — Auditoría de Compras No Comerciales"
TAREA = "Tarea 3.2: Revisión de aprobaciones de pedidos"
META = [("Auditor", "A. García"), ("Fecha", "12/05/2026"), ("Ref. Pentana", "CNC-2026-03/T3.2")]
PRUEBA = ("Se ha seleccionado una muestra de 45 pedidos de compra superiores a 10.000 € emitidos entre "
          "enero y marzo de 2026. Para cada pedido se ha verificado: (i) existencia de aprobación previa "
          "conforme a la matriz de delegación de facultades vigente; (ii) segregación entre solicitante y "
          "aprobador; (iii) existencia de tres ofertas comparativas cuando el importe supera 30.000 €.")
RESULTADOS = [
    ("R1", "En 6 de los 45 pedidos (13%) la aprobación se registró en el sistema con posterioridad a la fecha de emisión del pedido al proveedor.", "Anexo E1"),
    ("R2", "En 2 pedidos el aprobador coincidía con el solicitante, debido a una delegación temporal de permisos durante vacaciones que no fue revertida.", "Anexo E2"),
    ("R3", "En 3 de los 12 pedidos superiores a 30.000 € no consta la documentación de las tres ofertas comparativas en el expediente.", "Anexo E3"),
]
COMENTARIOS = ("El área indica que los 6 casos de aprobación posterior corresponden a pedidos urgentes de "
               "mantenimiento y que existe un correo previo del responsable, aunque fuera del sistema. "
               "Reconocen que la reversión de permisos tras vacaciones se gestiona manualmente.")
EVIDENCIAS = ["Extracto del módulo de compras (Anexo E1)", "Matriz de delegación de facultades v4.2 (Anexo E2)",
              "Expedientes de los pedidos de la muestra (Anexo E3)"]


def docx():
    import docx
    d = docx.Document()
    d.add_heading(TITULO, level=1)
    d.add_heading(TAREA, level=2)
    t = d.add_table(rows=0, cols=2)
    for k, v in META:
        fila = t.add_row().cells
        fila[0].text, fila[1].text = k, v
    d.add_heading("Prueba realizada", level=3)
    d.add_paragraph(PRUEBA)
    d.add_heading("Resultados", level=3)
    t = d.add_table(rows=1, cols=3)
    t.rows[0].cells[0].text, t.rows[0].cells[1].text, t.rows[0].cells[2].text = "Ref.", "Resultado", "Evidencia"
    for ref, txt, ev in RESULTADOS:
        c = t.add_row().cells
        c[0].text, c[1].text, c[2].text = ref, txt, ev
    d.add_heading("Comentarios del área auditada", level=3)
    d.add_paragraph(COMENTARIOS)
    d.add_heading("Evidencias", level=3)
    for e in EVIDENCIAS:
        d.add_paragraph(e, style="List Bullet")
    d.save(DESTINO / "papel_trabajo_compras.docx")


def xlsx():
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Tarea 3.2"
    ws.append([TITULO])
    ws.append([TAREA])
    ws.append([])
    for k, v in META:
        ws.append([k, v])
    ws.append([])
    ws.append(["Prueba realizada"])
    ws.append([PRUEBA])
    ws.append([])
    ws2 = wb.create_sheet("Resultados")
    ws2.append(["Ref.", "Resultado", "Evidencia"])
    for fila in RESULTADOS:
        ws2.append(list(fila))
    ws3 = wb.create_sheet("Comentarios y evidencias")
    ws3.append(["Comentarios del área auditada"])
    ws3.append([COMENTARIOS])
    ws3.append([])
    ws3.append(["Evidencias"])
    for e in EVIDENCIAS:
        ws3.append([e])
    wb.save(DESTINO / "papel_trabajo_compras.xlsx")


def pdf():
    """PDF mínimo escrito a mano (Helvetica, WinAnsi) con capa de texto real."""
    lineas = [TITULO, TAREA, "", " — ".join(f"{k}: {v}" for k, v in META), "", "Prueba realizada"]
    lineas += _envolver(PRUEBA) + ["", "Resultados"]
    for ref, txt, ev in RESULTADOS:
        lineas += _envolver(f"{ref}. {txt} ({ev})")
    lineas += ["", "Comentarios del área auditada"] + _envolver(COMENTARIOS) + ["", "Evidencias"]
    lineas += [f"- {e}" for e in EVIDENCIAS]

    def esc(s):
        return s.encode("cp1252", "replace").decode("cp1252").replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    contenido = "BT /F1 10 Tf 50 790 Td 13 TL\n" + "\n".join(f"({esc(l)}) Tj T*" for l in lineas) + "\nET"
    stream = contenido.encode("cp1252", "replace")
    objetos = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica /Encoding /WinAnsiEncoding >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objetos, 1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref = len(out)
    out += f"xref\n0 {len(objetos) + 1}\n0000000000 65535 f \n".encode()
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += f"trailer\n<< /Size {len(objetos) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode()
    (DESTINO / "papel_trabajo_compras.pdf").write_bytes(bytes(out))


def _envolver(texto, ancho=95):
    import textwrap
    return textwrap.wrap(texto, ancho)


if __name__ == "__main__":
    docx(); xlsx(); pdf()
    print(f"Generados en {DESTINO}: papel_trabajo_compras.docx / .xlsx / .pdf")
