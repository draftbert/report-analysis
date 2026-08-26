"""Extractor de DOCX con python-docx.

Recorre el documento en su orden real — párrafos y tablas intercalados tal
cual aparecen en el XML — en vez de `document.paragraphs` y `document.tables`
por separado, que rompería el orden y la estructura de secciones. Los
encabezados (`Heading 1..n`) se convierten en `Heading` con su nivel; el
resultado son las secciones que luego se ven en el Markdown final.

DOCX no tiene paginación real sin un motor de layout: la única frontera de
página fiable que existe en el propio fichero es el salto de página
explícito (manual o "salto antes de este párrafo"), así que es lo único que
usamos para delimitar `Page`.

Simplificación deliberada: no resolvemos el `numFmt` real de las listas
numeradas (exige leer numbering.xml); se decide ordenada/no-ordenada por el
nombre de estilo del párrafo. Suficiente para un PoC — no para reproducir
"a) / i) / 1)" tal cual.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from PIL import Image as PILImage

from .base import BaseExtractor
from .models import Block, Heading, Image, ListBlock, Page, Paragraph, Table
from .ocr import ocr_image


class DOCXExtractor(BaseExtractor):
    supported_formats = frozenset({"docx"})

    def extract_pages(self, path: Path) -> list[Page]:
        document = Document(str(path))
        builder = _PageBuilder()

        for item in _iter_block_items(document):
            if isinstance(item, DocxTable):
                builder.flush_list()
                builder.blocks.append(_table_block(item))
                continue

            builder.handle_paragraph(item, document)

        builder.flush_list()
        return builder.finish()


class _PageBuilder:
    def __init__(self) -> None:
        self.pages: list[Page] = []
        self.blocks: list[Block] = []
        self._list_items: list[str] = []
        self._list_ordered = False

    def handle_paragraph(self, paragraph: DocxParagraph, document: DocumentObject) -> None:
        if paragraph.paragraph_format.page_break_before:
            self._break_page()

        for image_bytes in _images_in_paragraph(paragraph, document):
            self.flush_list()
            self.blocks.append(_image_block(image_bytes))

        text = paragraph.text.strip()
        if text:
            level = _heading_level(paragraph)
            if level is not None:
                self.flush_list()
                self.blocks.append(Heading(text=text, level=level))
            elif _is_list_item(paragraph):
                self._add_list_item(text, ordered=_is_ordered_list(paragraph))
            else:
                self.flush_list()
                self.blocks.append(Paragraph(text=text))
        else:
            self.flush_list()

        if _has_manual_page_break(paragraph):
            self._break_page()

    def _add_list_item(self, text: str, *, ordered: bool) -> None:
        if self._list_items and self._list_ordered != ordered:
            self.flush_list()
        self._list_ordered = ordered
        self._list_items.append(text)

    def flush_list(self) -> None:
        if self._list_items:
            self.blocks.append(ListBlock(items=self._list_items, ordered=self._list_ordered))
            self._list_items = []

    def _break_page(self) -> None:
        self.flush_list()
        if self.blocks:
            self.pages.append(Page(number=len(self.pages) + 1, blocks=self.blocks))
            self.blocks = []

    def finish(self) -> list[Page]:
        if self.blocks or not self.pages:
            self.pages.append(Page(number=len(self.pages) + 1, blocks=self.blocks))
        return self.pages


def _iter_block_items(document: DocumentObject):
    """Recorrido en orden real del cuerpo del documento (receta estándar de
    python-docx): cada `<w:p>` como Paragraph, cada `<w:tbl>` como Table.
    """
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, document)


def _heading_level(paragraph: DocxParagraph) -> int | None:
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name == "Title":
        return 1
    if style_name.startswith("Heading"):
        digits = "".join(ch for ch in style_name if ch.isdigit())
        return int(digits) if digits else 1
    return None


def _is_list_item(paragraph: DocxParagraph) -> bool:
    style_name = paragraph.style.name if paragraph.style else ""
    if "List" in style_name:
        return True
    return paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None


def _is_ordered_list(paragraph: DocxParagraph) -> bool:
    style_name = paragraph.style.name if paragraph.style else ""
    return "Number" in style_name


def _images_in_paragraph(paragraph: DocxParagraph, document: DocumentObject) -> list[bytes]:
    """Bytes de cada imagen incrustada en los runs de este párrafo, en orden."""
    images = []
    for run in paragraph.runs:
        for blip in run._element.findall(f".//{qn('a:blip')}"):
            r_id = blip.get(qn("r:embed"))
            if r_id is None:
                continue
            part = document.part.related_parts.get(r_id)
            if part is not None:
                images.append(part.blob)
    return images


def _has_manual_page_break(paragraph: DocxParagraph) -> bool:
    for run in paragraph.runs:
        for br in run._element.findall(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
    return False


def _table_block(table: DocxTable) -> Table:
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    headers = rows[0] if rows else []
    body_rows = rows[1:]
    return Table(headers=headers, rows=body_rows)


def _image_block(image_bytes: bytes) -> Image:
    with PILImage.open(BytesIO(image_bytes)) as pil_image:
        ocr_text = ocr_image(pil_image.convert("RGB"))
    return Image(caption="Imagen incrustada", ocr_text=ocr_text)
